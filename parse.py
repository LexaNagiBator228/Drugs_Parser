import docx 
# pip install python-docx
import argparse
from difflib import SequenceMatcher

def similar(a, b):
    return SequenceMatcher(None, a, b).ratio()

def div_med(s):
    s = s.lower()
    dividers= [':','-','(','1','2','3','4','5','6','7','8','9','0','ампули','амп', 'ампулы']
    s = s.strip()
    s = s.strip('.')

    first_entrance = len(s) + 1 
    for el in dividers:
        entr = ' '+ el
        pos = s.find(entr)
        
        if(pos>0 and pos<first_entrance):
            first_entrance = pos
    
    if first_entrance == len(s) + 1 :
        name = s
        count = ''
    
    else:
        name = s[:first_entrance]
        count = s[first_entrance:].strip(' ,-.()').replace(')','').replace('(','')
    name = name.strip(' ,')
    other_names = name.split(',')
    return_list = []
    if(',' in name):
        for el in other_names:
            return_list.append([el.strip(' ,'),''])
        return_list[-1] = [other_names[-1].strip(' ,'), count]
        return return_list
    return [[name, count]]

    
def read_doc(path):
    doc = docx.Document(path)
    strings = [doc.paragraphs[i].text for i in range(len(doc.paragraphs))]
    return strings

def group_keys(keys):
    ## a bit greedy algo 
    group_of_keys = []
    if not keys:
        return []
    while keys:
        if not group_of_keys:
            group_of_keys.append(keys[0])
            keys.pop(0)
            continue
        # curr_el = keys[0]
        find_intersection = False
        for i in range(len(group_of_keys)):
            if(len(group_of_keys[i].intersection(keys[0]))>0):
                group_of_keys[i] = group_of_keys[i].union(keys[0])
                keys.pop(0)
                find_intersection = True
        if(find_intersection):
            continue
        group_of_keys.append(keys[0])
        keys.pop(0)
    group_of_keys = [list(el) for el in group_of_keys] 
    return group_of_keys    
        
def create_doc_file(dict_of_med):
    document = docx.Document()
    document.add_heading('Список ліків', 0)

    for key in sorted(dict_of_med, key= lambda x: dict_of_med[x]['Total number'], reverse=True):

        num  = dict_of_med[key]['Total number']
        document.add_heading(key + ' '+ str(num)+ ' разів', level=3)
        curr_spec = dict_of_med[key]['Spec']
        for med in  sorted(curr_spec, key=curr_spec.get, reverse=True):

            document.add_paragraph(
                        med + "\t\t, попросили "+ str(curr_spec[med])+' разів', style='List Bullet'
            )

    return document

def create_doc_file_short(dict_of_med):
    document = docx.Document()
    document.add_heading('Список ліків (короткий)', 0)

    for key in sorted(dict_of_med, key= lambda x: dict_of_med[x]['Total number'], reverse=True):

        num  = dict_of_med[key]['Total number']
        document.add_heading(key + ' '+ str(num)+ ' разів', level=3)


    return document
        


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument('--source_path', type=str, default='', help='path to source file, for example ./drugs.docx')
    parser.add_argument('--target_path', type=str, default='', help='path to file you want to save, for example ./res.docx')
    args = parser.parse_args() 
    path = args.source_path
    path_of_save = args.target_path
    if (path is '' or path_of_save is ''):
        path ='/home/user/personal/ліки.docx' 
        path_of_save = '/home/user/personal/res.docx'

    
    strings = read_doc(path)

    temp_list_of_med = list(map(div_med, strings))
    list_of_med = []
    for med in temp_list_of_med:
        for sub_med in med:
            list_of_med.append(sub_med)
            


    dict_of_med = dict()
    for el in list_of_med:
        if(el[0] in dict_of_med):
            dict_of_med[el[0]].append(el[1])
        else:
            dict_of_med[el[0]] = [el[1]]



    list_of_key = [el for el in sorted(dict_of_med)]

    similar_keys = []
    unique_keys = []
    for el1 in list_of_key:
        no_sim = True
        for el2 in list_of_key:
            score = similar(el1, el2)
            if(score>0.77 and not el1 is el2):  # hardcoded treshold 
                # print(el1, '\t', el2, score) 
                similar_keys.append(set([el1, el2]))
                no_sim = False
        if no_sim:
            unique_keys.append(el1)


    
    grouped_keys = group_keys(similar_keys)

    dict_of_unique_med = {el: [(el + ' '+sub_el).strip() for sub_el in dict_of_med[el]] for el in unique_keys  }

    # example of printing dict
    # for el in dict_of_unique_med:
    #     print()
    #     print(el)
    #     print(dict_of_unique_med[el])

    dict_of_grouped_med = {}
    for key_group in  grouped_keys:
        key = '/'.join(key_group)
        dict_of_grouped_med[key] = []
        for sub_key in key_group:
            list_for_sub_key = [(sub_key + ' '+sub_el).strip() for sub_el in dict_of_med[sub_key]] 
            dict_of_grouped_med[key] = list_for_sub_key + dict_of_grouped_med[key]


    final_dict = {**dict_of_unique_med, **dict_of_grouped_med}

    final_dict_with_spec ={}
    for key in final_dict:
        current_list = final_dict[key]
        count_dict = dict((x,current_list.count(x)) for x in set(current_list))
        total_count = 0
        for el in count_dict:
            total_count+= count_dict[el]
        final_dict_with_spec[key] = {'Spec':count_dict, 'Total number': total_count}

    doc = create_doc_file(final_dict_with_spec)
    doc.save(path_of_save)
    path_of_save_short = path_of_save.replace('.docx', '_short.docx') 
    doc_short = create_doc_file_short(final_dict_with_spec)
    doc_short.save(path_of_save_short)
